from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Optional
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from pydantic import BaseModel
from fastapi.responses import FileResponse
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml
from docx.oxml.ns import qn

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class CompanyDetails(BaseModel):
    company_name: str
    report_date: str
    prepared_by: str
    additional_details: str = ""
    description: str = ""
    cost_center: str = ""
    summary: str = ""
    references: str = ""
    impact: str = ""
    recommendation: str = ""

def parse_nessus_file(file_content: bytes) -> dict:
    """Parse Nessus XML file and extract vulnerability data."""
    try:
        root = ET.fromstring(file_content)
        vulnerabilities = {}
        
        # Find all ReportHost elements
        for host in root.findall(".//ReportHost"):
            ip_address = host.get('name')
            if ip_address not in vulnerabilities:
                vulnerabilities[ip_address] = []
            
            # Process each vulnerability
            for item in host.findall(".//ReportItem"):
                vuln = {
                    'plugin_id': item.get('pluginID'),
                    'plugin_name': item.get('pluginName'),
                    'severity': item.get('severity'),
                    'description': item.findtext('description', ''),
                    'solution': item.findtext('solution', ''),
                    'cvss_score': item.findtext('cvss_base_score', ''),
                }
                vulnerabilities[ip_address].append(vuln)
        
        return vulnerabilities
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error parsing Nessus file: {str(e)}")

def insert_vuln_table(doc, vulnerabilities):
    # Create a table with columns: Host, Plugin ID, Name, Severity, Description, Solution
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Host'
    hdr_cells[1].text = 'Plugin ID'
    hdr_cells[2].text = 'Name'
    hdr_cells[3].text = 'Severity'
    hdr_cells[4].text = 'Description'
    hdr_cells[5].text = 'Solution'
    for vuln in vulnerabilities:
        row_cells = table.add_row().cells
        row_cells[0].text = vuln['ip']
        row_cells[1].text = vuln['plugin_id']
        row_cells[2].text = vuln['plugin_name']
        row_cells[3].text = vuln['severity']
        row_cells[4].text = vuln['description']
        row_cells[5].text = vuln['solution']
    return table

def create_word_report(vulnerabilities: dict, company_details: BaseModel, logo_path: Optional[str] = None) -> Document:
    """Create a Word document with the specified structure."""
    doc = Document()

    # --- COVER PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Vulnerability Assessment Report")
    run.bold = True
    run.font.size = Pt(24)

    if logo_path:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(logo_path, width=Inches(2.0))
        except Exception:
            pass

    # Company Details Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.autofit = True
    table.cell(0, 0).text = "Company Name:"
    table.cell(0, 1).text = company_details.company_name
    table.cell(1, 0).text = "Date:"
    table.cell(1, 1).text = company_details.report_date
    table.cell(2, 0).text = "Prepared By:"
    table.cell(2, 1).text = company_details.prepared_by

    doc.add_page_break()

    # --- EXECUTIVE SUMMARY ---
    doc.add_heading("Executive Summary", level=1)
    
    # Count vulnerabilities by severity
    severity_counts = {'4': 0, '3': 0, '2': 0, '1': 0, '0': 0}
    for ip_vulns in vulnerabilities.values():
        for vuln in ip_vulns:
            severity = vuln.get('severity', '0')
            severity_counts[severity] += 1

    # Add summary text
    summary_text = f"""
This vulnerability assessment report presents the findings from the security scan conducted on {company_details.report_date}.
The scan identified {sum(severity_counts.values())} vulnerabilities across the target systems.

Severity Breakdown:
• Critical: {severity_counts['4']} vulnerabilities
• High: {severity_counts['3']} vulnerabilities
• Medium: {severity_counts['2']} vulnerabilities
• Low: {severity_counts['1']} vulnerabilities
• Informational: {severity_counts['0']} vulnerabilities

{company_details.description}
"""
    doc.add_paragraph(summary_text)

    # --- VULNERABILITY OVERVIEW TABLE (SINGLE TABLE, COLORED SEVERITY) ---
    doc.add_heading("Vulnerability Overview", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = [
        'Vulnerability Name',
        'Severity',
        'IP Address',
        'Description',
        'Impact',
        'Recommendation',
    ]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Severity color mapping
    severity_map = {'4': ('Critical', 'FF0000'),   # Red
                    '3': ('High', 'FFA500'),      # Orange
                    '2': ('Medium', 'FFFF00'),    # Yellow
                    '1': ('Low', '90EE90'),       # Light Green
                    '0': ('Info', 'ADD8E6')}      # Light Blue

    for ip, vulns in vulnerabilities.items():
        for vuln in vulns:
            row_cells = table.add_row().cells
            row_cells[0].text = vuln.get('plugin_name', '')
            severity = vuln.get('severity', '0')
            severity_label, color = severity_map.get(severity, ('Info', 'ADD8E6'))
            row_cells[1].text = severity_label
            # Color the severity cell
            shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{}"/>'.format(color))
            tc = row_cells[1]._tc
            tcPr = tc.get_or_add_tcPr()
            tcPr.append(shd)
            row_cells[2].text = ip
            row_cells[3].text = vuln.get('description', '')
            row_cells[4].text = str(vuln.get('cvss_score', ''))
            row_cells[5].text = vuln.get('solution', '')
    doc.add_paragraph()

    # --- DETAILED VULNERABILITY SECTIONS ---
    doc.add_page_break()
    doc.add_heading("Detailed Vulnerability Analysis", level=1)

    # Group vulnerabilities by severity for detailed section
    severity_order = [('4', 'Critical'), ('3', 'High'), ('2', 'Medium'), ('1', 'Low'), ('0', 'Info')]
    grouped_vulns = {sev: [] for sev, _ in severity_order}
    for ip, vulns in vulnerabilities.items():
        for vuln in vulns:
            severity = vuln.get('severity', '0')
            vuln_with_ip = vuln.copy()
            vuln_with_ip['ip'] = ip
            if severity in grouped_vulns:
                grouped_vulns[severity].append(vuln_with_ip)
            else:
                grouped_vulns['0'].append(vuln_with_ip)

    for severity, severity_label in severity_order:
        vulns = grouped_vulns[severity]
        if vulns:
            doc.add_heading(f"{severity_label} Vulnerabilities", level=2)
            for idx, vuln in enumerate(vulns, 1):
                doc.add_heading(f"{severity_label} Vulnerability {idx}", level=3)
                table = doc.add_table(rows=8, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                details = [
                    ('Vulnerability Name', vuln.get('plugin_name', '')),
                    ('CVSS', str(vuln.get('cvss_score', ''))),
                    ('OWASP Category', vuln.get('owasp_category', '')),
                    ('URL', vuln.get('url', '')),
                    ('Description', vuln.get('description', '')),
                    ('Impact', vuln.get('impact', '')),
                    ('Recommendation', vuln.get('solution', '')),
                    ('PoC', vuln.get('poc', '')),
                ]
                for i, (label, value) in enumerate(details):
                    table.cell(i, 0).text = label
                    table.cell(i, 1).text = value
                doc.add_paragraph()  # Add spacing between vulnerabilities

    return doc

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Handle multiple Nessus file uploads."""
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")
    
    all_vulnerabilities = {}
    
    for file in files:
        if not file.filename.endswith(('.nessus', '.xml')):
            raise HTTPException(status_code=400, detail=f"Invalid file format: {file.filename}")
        
        content = await file.read()
        vulnerabilities = parse_nessus_file(content)
        
        # Merge vulnerabilities
        for ip, vulns in vulnerabilities.items():
            if ip in all_vulnerabilities:
                all_vulnerabilities[ip].extend(vulns)
            else:
                all_vulnerabilities[ip] = vulns
    
    return {"message": "Files processed successfully", "host_count": len(all_vulnerabilities)}

@app.post("/generate-report")
async def generate_report(
    company_name: str = Form(...),
    report_date: str = Form(...),
    prepared_by: str = Form(...),
    additional_details: str = Form(""),
    report_filename: str = Form("vulnerability-report.docx"),
    description: str = Form(""),
    cost_center: str = Form(""),
    summary: str = Form(""),
    references: str = Form(""),
    impact: str = Form(""),
    recommendation: str = Form(""),
    logo: Optional[UploadFile] = File(None),
    files: List[UploadFile] = File(...)
):
    company_details = CompanyDetails(
        company_name=company_name,
        report_date=report_date,
        prepared_by=prepared_by,
        additional_details=additional_details,
        description=description,
        cost_center=cost_center,
        summary=summary,
        references=references,
        impact=impact,
        recommendation=recommendation
    )
    all_vulnerabilities = {}
    for file in files:
        content = await file.read()
        vulnerabilities = parse_nessus_file(content)
        for ip, vulns in vulnerabilities.items():
            if ip in all_vulnerabilities:
                all_vulnerabilities[ip].extend(vulns)
            else:
                all_vulnerabilities[ip] = vulns
    logo_path = None
    if logo is not None:
        logo_bytes = await logo.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(logo.filename)[-1]) as tmp_logo:
            tmp_logo.write(logo_bytes)
            logo_path = tmp_logo.name
    doc = create_word_report(all_vulnerabilities, company_details, logo_path=logo_path)
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    safe_filename = report_filename.strip() if report_filename.strip().endswith('.docx') else 'vulnerability-report.docx'
    return FileResponse(
        tmp_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=safe_filename
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

